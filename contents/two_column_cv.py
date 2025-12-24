from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

doc.add_heading('Your Full Name', 0)
doc.add_paragraph('Email | Phone | LinkedIn/GitHub | City, Country')

# Create a table for two-column layout
table = doc.add_table(rows=1, cols=2)
table.autofit = False
left_col = table.cell(0,0)
right_col = table.cell(0,1)

# Set widths (approx)
left_col.width = Inches(2.0)
right_col.width = Inches(4.0)

# Left Column Content
left_col_para = left_col.add_paragraph()
left_col_para.add_run('Skills\n').bold = True
left_col_para.add_run('ASP.NET MVC\nASP.NET Web Forms\nC# / .NET Framework\nSQL Server / T-SQL\nHTML, CSS, JavaScript\n\n')
left_col_para.add_run('Languages\n').bold = True
left_col_para.add_run('English – Fluent\nBengali – Native\n\n')
left_col_para.add_run('Tools / Platforms\n').bold = True
left_col_para.add_run('Visual Studio\nGit / GitHub\nPostman / Fiddler')

# Right Column Content
right_col_para = right_col.add_paragraph()
right_col_para.add_run('Professional Experience\n').bold = True
right_col_para.add_run('Software Engineer – Company Name, City, Country | Start Date – Present\n')
right_col_para.add_run('- Developed and maintained web applications using ASP.NET MVC and Web Forms.\n')
right_col_para.add_run('- Designed SQL Server databases and optimized queries for performance.\n')
right_col_para.add_run('- Implemented user authentication, role management, and data validation.\n\n')
right_col_para.add_run('Junior Developer – Previous Company Name, City, Country | Start Date – End Date\n')
right_col_para.add_run('- Assisted in building internal tools and web modules using ASP.NET.\n')
right_col_para.add_run('- Collaborated with team members to troubleshoot bugs and improve code quality.\n\n')
right_col_para.add_run('Education\n').bold = True
right_col_para.add_run('B.Sc. in Computer Science – Asian University of Bangladesh, Year of Graduation\n\n')
right_col_para.add_run('Projects / Highlights\n').bold = True
right_col_para.add_run('- E-commerce Web Application: Developed a fully functional e-commerce site using ASP.NET MVC and SQL Server.\n')
right_col_para.add_run('- Inventory Management System: Built a desktop/web hybrid system with real-time data management.')

# Save the document
doc_path = 'Two_Column_CV.docx'
doc.save(doc_path)
doc_path
